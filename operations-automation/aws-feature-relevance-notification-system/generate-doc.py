"""Generate the project overview Word document."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ============================================================
# TITLE PAGE
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('AWS Feature Relevance Notification System', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('SI Unified Operations Hackathon 2026')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x23, 0x2F, 0x3E)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Theme: Positioning (Value & ROI)').font.size = Pt(12)

doc.add_paragraph()
meta2 = doc.add_paragraph()
meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta2.add_run('Author: Rahul Srinivasan | Enterprise Pod').font.size = Pt(11)

doc.add_page_break()

# ============================================================
# EXECUTIVE SUMMARY
# ============================================================
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    'AWS publishes more than 200 feature announcements every month. For enterprise customers '
    'operating across dozens of AWS accounts, multiple business units, and hundreds '
    'of services — identifying which announcements matter to which workload is a manual, '
    'time-consuming process that often falls through the cracks.'
)
doc.add_paragraph(
    'The AWS Feature Relevance Notification System is an automated pipeline that ingests AWS '
    '"What\'s New" announcements, scores their relevance against customer workload profiles '
    'using Amazon Bedrock (Claude Sonnet 4), and delivers enriched Slack notifications with '
    'benefit categorization, impact estimates, and adoption recommendations — routed directly '
    'to the appropriate account team pod channel.'
)
doc.add_paragraph(
    'This system transforms 200+ monthly announcements into a curated, actionable feed '
    'per account team, ensuring no high-value feature goes unnoticed.'
)

# ============================================================
# BUSINESS PROBLEM
# ============================================================
doc.add_heading('2. The Business Problem', level=1)

doc.add_heading('The Pain Point', level=2)
doc.add_paragraph(
    'Today, TAMs manually scan AWS What\'s New, blogs, and release notes to identify features '
    'relevant to their customers. This approach has critical gaps:'
)
bullets = [
    'Volume overwhelm — 200+ announcements/month, impossible to read every one in depth',
    'Inconsistent coverage — depends on TAM bandwidth and familiarity with the customer\'s full architecture',
    'No workload context — a generic announcement about EKS means nothing without knowing the customer runs 340 Graviton3 nodes at $320K/month',
    'Missed value — features that could save $100K+/month go unnoticed because the TAM was focused on an incident that week',
    'No quantification — even when a feature is identified, articulating the dollar impact requires manual analysis',
]
for bullet in bullets:
    doc.add_paragraph(bullet, style='List Bullet')

doc.add_heading('The Impact', level=2)
doc.add_paragraph(
    'When features are missed or communicated without workload-specific context, customers '
    'miss cost savings, performance improvements, security enhancements, and simplification '
    'opportunities. This undermines the perceived value of UOps and makes renewal '
    'conversations harder — the customer can\'t see what they\'re getting.'
)

# ============================================================
# UOps HACKATHON RELEVANCE
# ============================================================
doc.add_heading('3. UOps Hackathon Relevance', level=1)

doc.add_paragraph(
    'This project directly addresses Theme 1: Positioning (Value & ROI) of the SI Unified '
    'Operations Hackathon 2026.'
)

doc.add_heading('Key Questions Answered', level=2)
questions = [
    'How might we quantify the dollar value UOps delivers and ground it in the customer\'s own workload context? → The system estimates impact in dollars, percentage savings, or performance gains specific to each workload.',
    'How might we demonstrate measurable improvement from a customer\'s own context? → Each notification includes workload-specific analysis ("your 340 Graviton3 nodes could save $150K/month").',
    'How might we link UOps value to the customer\'s overall workload cost and business criticality? → Tier-0 workloads get lower notification thresholds, ensuring mission-critical workloads never miss a relevant feature.',
]
for q in questions:
    doc.add_paragraph(q, style='List Bullet')

# ============================================================
# VALUE PROPOSITION
# ============================================================
doc.add_heading('4. Value for Customers and Account Teams', level=1)

doc.add_heading('For Customers', level=2)
customer_value = [
    'Never miss a feature that could save money, improve performance, or reduce risk',
    'Receive curated, workload-specific recommendations instead of generic announcements',
    'Get quantified impact estimates grounded in their actual architecture and scale',
    'Reduce time-to-adoption for high-value features from months to days',
]
for v in customer_value:
    doc.add_paragraph(v, style='List Bullet')

doc.add_heading('For Account Teams (TAMs/Pods)', level=2)
tam_value = [
    'Eliminate hours of manual What\'s New scanning per week',
    'Walk into customer conversations with pre-scored, workload-specific feature recommendations',
    'Strengthen renewal positioning with concrete, data-backed value delivery evidence',
    'Ensure consistent coverage even during high-incident weeks when proactive work gets deprioritized',
    'Scale expertise — junior TAMs get the same quality feature insights as senior architects',
]
for v in tam_value:
    doc.add_paragraph(v, style='List Bullet')

doc.add_heading('Quantified Value Example', level=2)
doc.add_paragraph(
    'For an enterprise customer with $750K/month in EKS compute spend, a single notification '
    'about Graviton4 support — scored at 91/100 relevance with an estimated $150K-$225K/month '
    'savings — delivers more value than a month of manual scanning. The system identified this '
    'opportunity in under 30 seconds with zero TAM effort.'
)

# ============================================================
# ARCHITECTURE
# ============================================================
doc.add_heading('5. Solution Architecture', level=1)

doc.add_paragraph(
    'The system follows an event-driven pipeline architecture with five core stages:'
)

stages = [
    ('Workload Onboarding', 'TAM creates a workload profile (workload name, key AWS services, free-text architecture description, account IDs, tier classification). Stored in DynamoDB.'),
    ('RSS Ingestion', 'EventBridge Scheduler triggers a Lambda function once per day to fetch and parse the AWS What\'s New RSS feed. Each announcement is deduplicated against a state table.'),
    ('Service Matching', 'The ingestion Lambda matches announcement keywords against each workload profile\'s key services list. Only matched announcements proceed to AI scoring.'),
    ('AI Relevance Scoring', 'AWS Step Functions orchestrates the scoring pipeline. A Lambda function passes the announcement + workload context to Amazon Bedrock (Claude Sonnet 4), which scores relevance (0-100) across 8 benefit categories.'),
    ('Notification Delivery', 'Announcements scoring ≥40 (or ≥25 for tier-0 workloads) are formatted into rich Slack Block Kit messages and delivered to the pod channel.'),
]
for i, (stage, desc) in enumerate(stages, 1):
    p = doc.add_paragraph()
    p.add_run(f'Stage {i}: {stage}').bold = True
    doc.add_paragraph(desc)

doc.add_heading('Architecture Diagram', level=2)
doc.add_picture(
    '/Users/rahulrsr/Documents/recipe-app/feature-relevance-poc/architecture.png',
    width=Inches(6.5)
)
doc.add_paragraph(
    'Figure 1: End-to-end pipeline from RSS ingestion through AI scoring to Slack delivery.',
    style='Caption'
)

# ============================================================
# HOW IT WORKS
# ============================================================
doc.add_heading('6. How It Works — Step by Step', level=1)

doc.add_heading('Daily Pipeline Flow', level=2)
steps = [
    'EventBridge Scheduler fires once per day → triggers RSS Ingestion Lambda.',
    'RSS Lambda fetches the AWS What\'s New RSS feed (~50 recent announcements).',
    'For each announcement, Lambda checks the DynamoDB state table — if already processed, skip.',
    'Lambda reads all workload profiles from DynamoDB and matches announcement keywords against each profile\'s key_services list.',
    'For each matched workload-announcement pair, Lambda triggers a Step Functions execution.',
    'Step Functions calls the Relevance Scorer Lambda, which combines the workload\'s free-text description with optional Knowledge Base context into a prompt.',
    'Amazon Bedrock (Claude Sonnet 4) scores relevance across 8 categories: cost optimization, performance, security, operational excellence, reliability, technical debt reduction, simplification, and new capability.',
    'Step Functions evaluates the threshold: score ≥ 40 proceeds to notification (≥ 25 for tier-0 workloads); below threshold is silently dropped.',
    'The Slack Notification Lambda formats the scoring output into a rich Block Kit message with benefit explanation, impact estimate, effort level, and risk assessment.',
    'Notification is delivered to the configured Slack channel via webhook (URL stored in Secrets Manager).',
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_heading('Workload Onboarding (One-Time Setup)', level=2)
doc.add_paragraph(
    'A TAM onboards a workload by providing:'
)
onboard_fields = [
    'Workload name and ID',
    'AWS account IDs (supports multiple accounts per workload)',
    'Key AWS services used (for keyword matching)',
    'Free-text description (architecture, scale, constraints, pain points — this is what Claude reads for deep scoring)',
    'Workload tier (tier-0 for mission-critical, standard for everything else)',
    'Pod Slack channel for notification routing',
]
for f in onboard_fields:
    doc.add_paragraph(f, style='List Bullet')

doc.add_paragraph(
    'Optionally, architecture documents can be uploaded to a Bedrock Knowledge Base for '
    'even deeper contextual scoring.'
)

# ============================================================
# BENEFIT CATEGORIES
# ============================================================
doc.add_heading('7. Benefit Categories', level=1)
doc.add_paragraph(
    'The system scores relevance across 8 dimensions, ensuring value is captured beyond just cost:'
)

table = doc.add_table(rows=9, cols=3)
table.style = 'Light Grid Accent 1'
headers = ['Category', 'Tag', 'What It Captures']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h

categories = [
    ('Cost Optimization', '💰', 'Reduces spend, better pricing, savings plans'),
    ('Performance', '⚡', 'Lower latency, higher throughput, capacity'),
    ('Security & Compliance', '🔒', 'New controls, compliance, encryption'),
    ('Operational Excellence', '🔧', 'Simplifies management, observability'),
    ('Reliability', '🛡️', 'Higher availability, better DR'),
    ('Technical Debt Reduction', '🧹', 'Deprecation replacements, modernization'),
    ('Simplification', '✨', 'Replaces complex workarounds'),
    ('New Capability', '🚀', 'Unlocks previously impossible patterns'),
]
for i, (cat, tag, desc) in enumerate(categories, 1):
    table.rows[i].cells[0].text = cat
    table.rows[i].cells[1].text = tag
    table.rows[i].cells[2].text = desc

# ============================================================
# TECHNOLOGY STACK
# ============================================================
doc.add_heading('8. Technology Stack', level=1)

tech_table = doc.add_table(rows=8, cols=3)
tech_table.style = 'Light Grid Accent 1'
tech_headers = ['AWS Service', 'Role', 'Why This Service']
for i, h in enumerate(tech_headers):
    tech_table.rows[0].cells[i].text = h

tech = [
    ('EventBridge Scheduler', 'Daily trigger', 'Serverless cron, no infrastructure to manage'),
    ('AWS Lambda', 'Compute (4 functions)', 'Event-driven, pay-per-invocation, scales to zero'),
    ('AWS Step Functions', 'Orchestration', 'Visual debugging, built-in retries, threshold gates'),
    ('Amazon DynamoDB', 'State management', 'Serverless, pay-per-request, TTL for auto-cleanup'),
    ('Amazon Bedrock (Claude)', 'AI scoring & analysis', 'Managed foundation models, no ML ops'),
    ('Bedrock Knowledge Bases', 'Document retrieval (optional)', 'Semantic search over architecture docs'),
    ('Secrets Manager', 'Credential storage', 'Secure webhook URL storage, rotation support'),
]
for i, (svc, role, why) in enumerate(tech, 1):
    tech_table.rows[i].cells[0].text = svc
    tech_table.rows[i].cells[1].text = role
    tech_table.rows[i].cells[2].text = why

# ============================================================
# DEMO RESULTS
# ============================================================
doc.add_heading('9. Demo Results', level=1)
doc.add_paragraph(
    'In live testing with a simulated EKS Graviton4 announcement scored against a '
    'Digital Experience Platform workload profile:'
)

results = [
    ('Relevance Score', '91/100'),
    ('Primary Category', '💰 Cost Optimization'),
    ('Secondary Categories', '⚡ Performance, 🔧 Operational Excellence'),
    ('Estimated Impact', '~$150K-$225K/month savings (20-30% of $750K compute)'),
    ('Effort to Adopt', 'Configuration change (managed node group update)'),
    ('Time from Announcement to Notification', '< 30 seconds'),
]
for label, value in results:
    p = doc.add_paragraph()
    p.add_run(f'{label}: ').bold = True
    p.add_run(value)

# ============================================================
# FUTURE ENHANCEMENTS
# ============================================================
doc.add_heading('10. Future Enhancements (Path to Production)', level=1)
enhancements = [
    'API Gateway + web UI for self-service workload onboarding',
    'Cost Explorer integration for automatic service discovery per account',
    'Bedrock Knowledge Base with full architecture documents for deeper scoring',
    'Historical tracking — dashboard showing features recommended, adopted, and value realized',
    'Email/Teams/Chime delivery options beyond Slack',
    'Feedback loop — TAMs can thumbs-up/down notifications to improve scoring over time',
    'Multi-customer support with pod-based routing across entire SI portfolio',
]
for e in enhancements:
    doc.add_paragraph(e, style='List Bullet')

# Save
output_path = '/Users/rahulrsr/Documents/recipe-app/feature-relevance-poc/AWS_Feature_Relevance_System_Overview.docx'
doc.save(output_path)
print(f"Document saved to: {output_path}")
